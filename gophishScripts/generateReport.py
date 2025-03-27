import requests
import argparse
import json
from docx import Document
import urllib3
from datetime import datetime
import getModlishkaCreds
import sys

#Suppress InsecureRequestWarnings
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)


# Function to make API requests
def make_request(endpoint):
    url = f"{SERVER_URL}{endpoint}"
    headers = {"Authorization": API_KEY}
    
    try:
        response = requests.get(url, headers=headers, verify=False)
        response.raise_for_status()
        return response.json()
    except requests.exceptions.RequestException as e:
        print(f"Error: {e}")
        return None


def generate_word_report(data, campaign_name):
    output_file = f"{args.client} Phishing Campaign Report.docx"
    try:
        doc = Document(output_file)
    except Exception:
        doc = Document()

    heading = doc.add_heading(f"Phishing Results for Campaign {campaign_name}", level=1)
    heading.alignment = 1
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    
    headers = ["Email", "Opened", "Clicked", "Data Submitted", "Session","Date Modified"]
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].alignment = 1
    
    for record in data:
        email = record.get("email")
        if record.get("status") == "Submitted Data":
            opened = "Yes"
            clicked = "Yes"
            submitted_data = "Yes"
        elif record.get("status") == "Clicked Link":
            opened = "Yes"
            clicked = "Yes"
            submitted_data = "No"
        else:
            opened = "Yes"
            clicked = "No"
            submitted_data = "No" 
        terminate = record.get("terminate")
        if terminate == "Y":
            terminate = "Yes"
        else:
            terminate = "No" 
        modDate = str(datetime.strptime(record.get("modified_date")[:19], "%Y-%m-%dT%H:%M:%S"))

        row_cells = table.add_row().cells
        row_cells[0].text = email
        row_cells[1].text = opened
        row_cells[2].text = clicked
        row_cells[3].text = submitted_data
        row_cells[4].text = terminate
        row_cells[5].text = modDate
    
        for cell in row_cells:
            cell.paragraphs[0].alignment = 1

    doc.save(output_file)
    print(f"Report saved to {output_file}")

def parse_file(file_path): 
    try:
        with open(file_path, 'r') as file:
            data = file.read()
    except FileNotFoundError:
        print(f"Error: The file '{file_path}' was not found.")
        sys.exit(1)

    records = []
    records_data = data.strip().split('\n\n')
    
    # Loop through each record and extract the key-value pairs
    for record in records_data:
        # Split each record into lines
        lines = record.split('\n')
        
        # Initialize an empty dictionary for each record
        record_dict = {}

        for line in lines:
            # Split the line into key and value
            key, value = line.split(': ', 1)
            record_dict[key] = value
        
        # Append the dictionary to the records list
        records.append(record_dict)
    
    return records



if __name__ == "__main__":

    parser = argparse.ArgumentParser(description="Generate a Word Doc Gophish Report.")
    parser.add_argument("--db", help="Name of the Modlishka Database")
    parser.add_argument("--client", help="Name of the Client")
    args = parser.parse_args()

    # Prompt for API key
    API_KEY = input("Enter your GoPhish API key: ")

    # GoPhish server URL
    SERVER_URL = "http://localhost:3333"

    # Validate inputs
    if not API_KEY or not SERVER_URL:
        print("API key and server URL are required.")
        exit(1)

    #uuidList = getModlishkaCreds.read_redis_db(args.db)
    uuidList = parse_file(args.db)
    print(f"uuid list is {uuidList}")

    campaigns = make_request("/api/campaigns")

    while True: 
        # Display campaign list
        print("\nAvailable Campaigns:")
        for campaign in campaigns:
            print(f"ID: {campaign['id']}, Name: {campaign['name']}")

        # Prompt for campaign ID
        campaign_Numbers = input("Enter the ID/IDs (comma separated) of the campaign you want to export or type 'done': ")

        if campaign_Numbers == "done":
            print("Done generating reports. Goodbye!")
            break
        
        campaignIDList = campaign_Numbers.split(',')
        for campaign_id in campaignIDList: 

            # Export campaign results
            print(f"Exporting results and generating report for campaign ID: {campaign_id}...")

            results = make_request(f"/api/campaigns/{campaign_id}/results")


            if not results:
                print("Failed to export campaign results. Ensure the campaign ID is correct.")
                continue

            #Retrieving Campaign Name 
            campaign_name = results.get("name", f"campaign_{campaign_id}")
            campaign_name = campaign_name.replace(" ", "_")

            filtered_results = []
            for result in results['results']:
                if result['status'] in ['Clicked Link', 'Email Opened', 'Submitted Data']:
                    # Extract the desired fields
                    status = result['status']
                    terminationStatus = "N"
                    for uuidRecord in uuidList:
                        
                        if result['id'] == uuidRecord['UUID']:
                            print(f"user {uuidRecord['Username']} submitted data! adding to report. They're in trouble!")
                            status = 'Submitted Data'
                            terminationStatus = uuidRecord['Terminated']
                            continue
                      
                            
                    extracted_data = {
                        'email': result['email'],
                        'status': status,
                        'terminate': terminationStatus,
                        'send_date': result['send_date'],
                        'modified_date': result['modified_date']
                    }
                    filtered_results.append(extracted_data)
            print(f"filtered results is {filtered_results}")
            
            generate_word_report(filtered_results, campaign_name)
